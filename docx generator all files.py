import pandas as pd
import docx
from tqdm import tqdm
from time import sleep
from colorama import init
from termcolor import colored
# importing mapping files
import country_mappings1 as maps1
import country_mappings2 as maps2
import country_mappings3 as maps3
import country_mappings4 as maps4
import country_mappings5 as maps5

# getting classes as getmap
text1 = 'Customer attributes:'
text2 = 'Program attributes:'
text3 = 'Lineitem_t attributes:'
text4 = 'Lineitem_r attributes:'
text5 = 'Lineitem_dol attributes:'

# getting the path of attributes in files
getmap1 = maps1.maps_az_con()
getmap2 = maps2.maps_az_srv()
getmap3 = maps3.maps_br_sac()
getmap4 = maps4.maps_nl_srv()
getmap5 = maps5.maps_ui_srv()

# list of elements needed for the final loop (getting attributes)
getmap = [getmap1, getmap2, getmap3, getmap4, getmap5]

# function for creating tables from dataframes
def save_doc(df, text):
    doc.add_paragraph('')
    doc.add_paragraph(f'{text}') # adding text in front of the table
    t = doc.add_table(df.shape[0]+1, df.shape[1], style="Table Grid") # adding table with extra row for headers
    # adding header row
    for j in range(df.shape[-1]):
        t.cell(0,j).text = df.columns[j]
    # adding the rest of the data frame
    for i in range(df.shape[0]):
        for j in range(df.shape[-1]):
            t.cell(i+1,j).text = str(df.values[i,j])


# names of files needed for saving docx files
name = ['country_mappings1', 'country_mappings2', 'country_mappings3', 
'country_mappings4', 'country_mappings5']

# loop that generates files for each country_mapping file
for i in tqdm(range(5), desc = colored("Update on the files genarator:", "red"), colour="green"):
    sleep(0.0000005)
    doc = docx.Document()
    # saving customer and program tables in docx file
    customer = pd.DataFrame(getmap[i].header_map.items(), columns=['Xml file', 'Excel file'])
    program = pd.DataFrame(getmap[i].program_map.items(), columns=['Xml file', 'Excel file'])
    save_doc(customer, text1)
    save_doc(program, text2)
    # for each file lineitem components are different- if statements needed for getting the right ones
    # 'country_mappings1' file's lineitem component
    if i == 0:
        lineitem_tiered = pd.DataFrame(getmap[i].lineitem_tieredli_map.items(), columns=['Xml file', 'Excel file'])
        save_doc(lineitem_tiered, text3)
    # 'country_mappings2', 'country_mappings3' files' lineitem components
    elif i == 1 or i == 2:
        lineitem_rental = pd.DataFrame(getmap[i].lineitem_rental_map.items(), columns=['Xml file', 'Excel file'])
        lineitem_tiered = pd.DataFrame(getmap[i].lineitem_tieredli_map.items(), columns=['Xml file', 'Excel file'])
        lineitem_dol = pd.DataFrame(getmap[i].lineitem_dol_map.items(), columns=['Xml file', 'Excel file'])
        save_doc(lineitem_rental, text4)
        save_doc(lineitem_tiered, text3)
        save_doc(lineitem_dol, text5)
    # 'country_mappings4', 'country_mappings5' files' lineitem components
    elif i == 3 or i == 4:
        lineitem_rental = pd.DataFrame(getmap[i].lineitem_rental_map.items(), columns=['Xml file', 'Excel file'])
        save_doc(lineitem_rental, text4)
    # saving tables as a file for each country_mapping file
    doc.save(f'output {name[i]}.docx')

print(colored("Docx files exported", "green"))




 